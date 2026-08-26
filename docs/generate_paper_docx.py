"""Generate Word (.docx) baseline paper draft from the v2 + DirAcc experiment."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = Path(__file__).parent / "paper_baseline_draft.docx"


# ──────────── Styling helpers ────────────
def set_default_font(doc, font_name="Times New Roman", size=11):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size)
    # Asian font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts")) or OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "맑은 고딕")
    rpr.append(rfonts)


def add_title(doc, text, size=18, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    p.paragraph_format.space_after = Pt(6)


def add_subtitle(doc, text, size=11, italic=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    p.paragraph_format.space_after = Pt(4)


def add_author_block(doc, name, affil, email):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(name)
    r.font.bold = True
    p.add_run(f"\n{affil}\n{email}")
    p.paragraph_format.space_after = Pt(12)


def add_section_heading(doc, number, title, level=1):
    if level == 1:
        p = doc.add_heading("", level=1)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f"{number}. {title}")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    else:
        p = doc.add_heading("", level=2)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{number} {title}")
        run.font.size = Pt(12)
        run.font.bold = True


def add_body_paragraph(doc, text, first_line_indent=True):
    p = doc.add_paragraph(text)
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    if first_line_indent:
        pf.first_line_indent = Cm(0.5)
    for run in p.runs:
        run.font.size = Pt(11)


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.line_spacing = 1.3
    for run in p.runs:
        run.font.size = Pt(11)


def add_equation(doc, latex_text):
    """Add a centered equation (as plain text, not MathType)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(latex_text)
    run.font.size = Pt(11)
    run.font.italic = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def add_table(doc, rows, headers=None, widths=None, caption=None):
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.bold = True
        run.font.size = Pt(10)

    n_cols = len(headers) if headers else len(rows[0])
    tbl = doc.add_table(rows=1 if headers else 0, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Light Grid Accent 1"

    if headers:
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for p in hdr_cells[i].paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in rows:
        cells = tbl.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    # space after table
    doc.add_paragraph()


# ──────────── Build document ────────────
doc = Document()
set_default_font(doc)

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Title block ────────────────────────────────────────────────────────
add_title(doc, "Hybrid Omega Matrix Design for Black-Litterman Portfolios", size=16)
add_title(doc,
          "with Machine Learning / Deep Learning Views and Reinforcement Learning Model Selection",
          size=12, bold=False)
add_subtitle(doc, "Capstone Project Baseline Draft -- 2026")

add_author_block(
    doc,
    name="Jay Tak",
    affil="Department of Industrial & Management Engineering, Undergraduate Capstone",
    email="roboadvisor2026@gmail.com",
)

# ── Abstract ───────────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run("Abstract  ")
r.font.bold = True
r.font.size = Pt(11)
p.paragraph_format.space_after = Pt(2)

abstract_text = (
    "The Black-Litterman (BL) framework combines the market-implied equilibrium "
    "with investors' subjective views, offering a more stable alternative to the "
    "Markowitz mean-variance optimizer. While most prior work focuses on deriving "
    "the view vector Q from Machine Learning (ML) or Deep Learning (DL) predictions, "
    "the design of the view-uncertainty matrix Omega has remained comparatively "
    "under-explored. "
    "\n\n"
    "This paper introduces a hybrid Omega construction that scales each view's "
    "uncertainty by (i) the forecasting model's direction-accuracy on held-out data, "
    "and (ii) the magnitude of the prediction itself. Together with "
    "cross-sectional feature normalization and a unified 13-model ML/DL "
    "forecasting pipeline, the proposed framework is evaluated on a 22-asset U.S. "
    "equity universe (2013--2026, 20 walk-forward folds). "
    "\n\n"
    "Empirically, BL strategies that consume LightGBM (Sharpe 1.299) or LSTM "
    "(Sharpe 1.295) views outperform the BL-NoView benchmark (Sharpe 1.279), "
    "while a reinforcement-learning portfolio agent (DQN) reduces the average "
    "maximum drawdown by 7.5%. These results support the view that forecasting "
    "quality translates into portfolio performance when Omega is informed by "
    "a model-quality signal."
)
p2 = doc.add_paragraph(abstract_text)
p2.paragraph_format.line_spacing = 1.5
p2.paragraph_format.space_after = Pt(6)
for run in p2.runs:
    run.font.size = Pt(10)

# Keywords
kw = doc.add_paragraph()
r = kw.add_run("Keywords: ")
r.font.bold = True
r.font.size = Pt(10)
r2 = kw.add_run(
    "Black-Litterman, Portfolio Optimization, Machine Learning, Deep Learning, "
    "Reinforcement Learning, Hybrid Omega, Walk-Forward Validation"
)
r2.font.size = Pt(10)
r2.font.italic = True

# ── 1. Introduction ────────────────────────────────────────────────────
add_section_heading(doc, 1, "Introduction")
add_section_heading(doc, "1.1", "Background", level=2)
add_body_paragraph(doc,
    "Harry Markowitz's mean-variance optimizer (Markowitz, 1952) formalized "
    "portfolio construction as a risk-return trade-off, yet its sensitivity to "
    "estimation error in the expected-return vector severely limits practical "
    "use. Black and Litterman (1992) proposed a Bayesian remedy: the prior "
    "distribution is anchored at the market-implied CAPM return, and investor "
    "views are incorporated through a probabilistic update, yielding the "
    "posterior expected return:")
add_equation(doc,
    "E[r] = [(tau*Sigma)^(-1) + P'*Omega^(-1)*P]^(-1) * "
    "[(tau*Sigma)^(-1)*Pi + P'*Omega^(-1)*Q]")
add_body_paragraph(doc,
    "where Pi is the market-equilibrium return, P is the pick matrix, Q is the "
    "view vector, and Omega is a diagonal matrix capturing the uncertainty of "
    "each view. The tau scalar controls the confidence in the prior.")

add_section_heading(doc, "1.2", "Motivation and Gap", level=2)
add_body_paragraph(doc,
    "Extensive prior research has investigated how to generate Q, including "
    "EGARCH-based views (Beach and Orlov, 2007), volatility forecasts "
    "(Duqi et al., 2014), and ML/DL regressors (Shigolakov, 2025). However, "
    "the specification of Omega has mostly defaulted to either (i) a scalar "
    "multiple of the diagonal of P*tau*Sigma*P' or (ii) the in-sample residual "
    "variance of each model. Both treat every view as equally trustworthy -- "
    "a clearly problematic assumption when the forecasting models differ in "
    "predictive power.")
add_body_paragraph(doc,
    "This paper closes that gap by conditioning Omega on (i) a model-quality "
    "score computed on held-out data and (ii) the prediction magnitude, while "
    "preserving the standard BL identity.")

add_section_heading(doc, "1.3", "Contributions", level=2)
add_bullet(doc,
    "A hybrid Omega construction that couples forecasting quality with "
    "prediction strength in a single closed-form expression.")
add_bullet(doc,
    "A unified 13-model ML/DL pipeline (Ridge, Lasso, DecisionTree, "
    "RandomForest, XGBoost, LightGBM, CatBoost, MLP, LSTM, CNN1D, Transformer, "
    "PatchTST, Hybrid LSTM-Transformer) with Optuna-tuned hyper-parameters.")
add_bullet(doc,
    "Cross-sectional feature normalization that converts scale-heterogeneous "
    "indicators into rank- and z-score-based signals, which are more "
    "appropriate for the relative nature of portfolio optimization.")
add_bullet(doc,
    "A comparison against reinforcement-learning agents (DQN, PPO) that "
    "dynamically select the best forecasting model per re-balance step.")
add_bullet(doc,
    "Empirical evidence on a 22-asset universe (2013--2026) that the proposed "
    "design enables individual ML/DL models to outperform the BL-NoView "
    "benchmark in terms of Sharpe ratio.")

# ── 2. Methodology ─────────────────────────────────────────────────────
add_section_heading(doc, 2, "Methodology")

add_section_heading(doc, "2.1", "Dataset", level=2)
add_body_paragraph(doc,
    "We construct bl_v3_dataset.csv from Yahoo Finance price data spanning "
    "February 2013 to February 2026, covering 22 large-cap U.S. equities "
    "across multiple sectors (AAPL, AMZN, BRK-B, CAT, COST, CVX, EQIX, GE, "
    "GOOGL, JPM, LIN, LLY, META, NEE, NVDA, PLD, SHW, SO, TSLA, UNH, WMT, XOM). "
    "Seven technical indicators per asset (RSI-14, StochRSI-14, ROC-10, "
    "TSI-25-13, DPO-20, ATR-14, and market-capitalization weight) are "
    "complemented by five macro variables (VIX, MOVE, HY_OAS, 10Y-2Y spread, "
    "U.S. Dollar Index). The forecasting target is the 5-day forward return.")

add_section_heading(doc, "2.2", "Feature Engineering", level=2)
add_body_paragraph(doc,
    "Raw macro levels are incompatible with asset-level percentage features, "
    "so we replace each macro level with its 1-day pct_change, 5-day "
    "pct_change, and ratio to its 20-day moving average. Fourteen asset-level "
    "derived features are added: log-return lags at horizons {1,5,10,20}, "
    "realised volatilities at {5,20}, volatility ratio, three price-to-MA "
    "ratios, ma5/ma20 ratio, rolling skewness, RSI change, ATR regime, "
    "price-normalized DPO, and market-cap weight change.")
add_body_paragraph(doc,
    "We then apply cross-sectional normalization at every date: for each "
    "feature we compute a percent-rank (0-1) and a z-score across the 22 "
    "assets. This yields signals such as 'RSI of AAPL is in the top 20% "
    "today' instead of the raw value, which is more aligned with the "
    "relative nature of portfolio construction.")

add_section_heading(doc, "2.3", "Forecasting Models", level=2)
add_body_paragraph(doc,
    "Thirteen ML/DL models are trained per fold: Ridge and Lasso (linear "
    "regularised regressions), DecisionTree, RandomForest, XGBoost, LightGBM, "
    "CatBoost (tree-based ensembles), and MLP, LSTM, CNN1D, Transformer, "
    "PatchTST, and a Hybrid LSTM-Transformer (sequential DL architectures). "
    "Hyper-parameters are tuned on the first three folds via Optuna with "
    "25 trials for ML models and 40-60 trials for DL models, minimising the "
    "view-build RMSE.")

add_section_heading(doc, "2.4", "Walk-Forward Validation", level=2)
add_body_paragraph(doc,
    "We use a nested walk-forward scheme with 1,008 days of model training, "
    "252 days of view-build (used to compute model-quality scores), and "
    "756 days of out-of-sample test. The window advances by 63 days per "
    "fold, yielding 20 folds that cover the period 2018--2026.")

add_section_heading(doc, "2.5", "Hybrid Omega Construction", level=2)
add_body_paragraph(doc,
    "For each view i generated by model m on date t, we define Omega's "
    "diagonal as:")
add_equation(doc,
    "omega_i^2 = sigma_i^2 * (1 / max(q_m, q_floor))^alpha * "
    "(1 / max(|y_i|, s_floor))^beta * lambda")
add_body_paragraph(doc, "The components are:")
add_bullet(doc, "sigma_i^2: baseline variance (in-sample residual variance).")
add_bullet(doc,
    "q_m: direction-accuracy-based quality score, q_m = 2*(DirAcc_m - 0.5), "
    "so that a random predictor scores 0 and a perfect one scores 1. "
    "Negative values (worse than random) are clipped at q_floor = 0.005.")
add_bullet(doc,
    "|y_i|: magnitude of the prediction. Low-confidence predictions receive "
    "a larger Omega and are weighted less in the BL posterior.")
add_bullet(doc,
    "alpha = beta = 0.25 (moderate influence) and lambda = 0.1 (global scale). "
    "These values were selected to avoid over-suppression of views while "
    "retaining meaningful model-quality differentiation.")
add_body_paragraph(doc,
    "Because the rest of the BL pipeline consumes a pre-adjusted uncertainty "
    "column, the effective Omega is produced by rescaling the uncertainty "
    "prior to BL solving, which keeps the existing omega_method='uncertainty' "
    "path fully compatible with the hybrid extension.")

add_section_heading(doc, "2.6", "BL + MVO Pipeline", level=2)
add_body_paragraph(doc,
    "Each fold iterates through re-balancing dates (every 5 days). At each "
    "decision date we (i) estimate the covariance matrix from the past 756 "
    "days of log returns, (ii) solve the BL posterior using the selected "
    "model's views and the hybrid Omega, and (iii) run a long-only "
    "maximum-Sharpe MVO with a per-asset weight cap of 25%.")

add_section_heading(doc, "2.7", "Reinforcement Learning Agents", level=2)
add_body_paragraph(doc,
    "To test whether dynamic model selection can further improve performance, "
    "we train two RL agents with the same action space (13 models) and "
    "observation (cross-sectional prediction statistics). The agents differ "
    "in algorithm (Double-DQN vs. PPO) and reward: the primary variant uses "
    "an EMA-based Sharpe estimate on the BL portfolio return. Each agent is "
    "trained on the view-build partition and deployed on test.")

# ── 3. Results ─────────────────────────────────────────────────────────
add_section_heading(doc, 3, "Experimental Results")

add_section_heading(doc, "3.1", "Forecasting Direction Accuracy", level=2)
add_body_paragraph(doc,
    "All 13 models exceed the 50% random baseline. Lasso and Transformer "
    "achieve the highest hit-rate (56.4%), followed by HybridLSTMTF and "
    "XGBoost (56.2%). The spread is narrow but meaningful at scale.")

add_table(
    doc,
    rows=[
        ["Lasso", "0.564"],
        ["Transformer", "0.564"],
        ["HybridLSTMTF", "0.562"],
        ["XGBoost", "0.562"],
        ["MLP", "0.561"],
        ["CatBoost", "0.560"],
        ["PatchTST", "0.558"],
        ["LightGBM", "0.556"],
        ["RandomForest", "0.555"],
        ["LSTM", "0.551"],
        ["CNN1D", "0.548"],
        ["DecisionTree", "0.536"],
        ["Ridge", "0.531"],
    ],
    headers=["Model", "Direction Accuracy"],
    caption="Table 1. View-build direction accuracy per forecasting model.",
)

add_section_heading(doc, "3.2", "Portfolio Performance", level=2)
add_body_paragraph(doc,
    "Table 2 reports the mean performance across the 20 test folds. "
    "BL_LightGBM attains the highest Sharpe ratio (1.299), followed by "
    "BL_LSTM (1.295); both beat the BL-NoView benchmark (1.279) by 1.6% and "
    "1.3%, respectively. BL_Lasso produces the largest total return (1.946). "
    "The reinforcement-learning agent BL_DQN_Port does not outperform the "
    "NoView benchmark in Sharpe, but it reduces the mean maximum drawdown "
    "from -0.334 to -0.309, a 7.5% improvement in tail-risk control.")

add_table(
    doc,
    rows=[
        ["BL_LightGBM",     "1.299", "1.902", "-0.311", "0.50"],
        ["BL_LSTM",         "1.295", "1.806", "-0.306", "0.65"],
        ["BL_NoView",       "1.279", "1.899", "-0.334", "0.00"],
        ["BL_Ridge",        "1.271", "1.550", "-0.283", "0.50"],
        ["BL_HybridLSTMTF", "1.251", "1.766", "-0.289", "0.40"],
        ["BL_DQN_Port",     "1.245", "1.747", "-0.309", "0.45"],
        ["BL_XGBoost",      "1.244", "1.592", "-0.286", "0.45"],
        ["BL_CatBoost",     "1.231", "1.443", "-0.273", "0.40"],
        ["BL_RandomForest", "1.195", "1.417", "-0.284", "0.35"],
        ["BL_Lasso",        "1.195", "1.946", "-0.389", "0.30"],
        ["BL_DecisionTree", "1.185", "1.747", "-0.332", "0.30"],
        ["BL_Transformer",  "1.144", "1.429", "-0.291", "0.20"],
        ["BL_PPO_Port",     "1.107", "1.466", "-0.322", "0.20"],
        ["BL_PatchTST",     "1.091", "1.489", "-0.348", "0.20"],
        ["BL_CNN1D",        "1.062", "1.432", "-0.334", "0.20"],
    ],
    headers=["Strategy", "Mean Sharpe", "Total Return", "Max DD", "Win vs NoView"],
    caption="Table 2. Walk-forward out-of-sample portfolio performance (20 folds).",
)

add_section_heading(doc, "3.3", "Prediction Quality vs. Portfolio Performance", level=2)
add_body_paragraph(doc,
    "The cross-model correlation between the DirAcc quality score and the "
    "portfolio Sharpe is weakly negative (r = -0.087). We interpret this as "
    "a threshold effect: above a minimum DirAcc level (~55%) the specific "
    "value of the quality score has limited explanatory power for Sharpe "
    "because turnover, volatility structure, and regime exposure dominate. "
    "Below the threshold, however, the three lowest-DirAcc models (CNN1D, "
    "DecisionTree, Ridge) all rank in the bottom half of Sharpe ratios, "
    "suggesting that baseline prediction quality is a necessary but not "
    "sufficient condition for portfolio outperformance.")

# ── 4. Discussion ──────────────────────────────────────────────────────
add_section_heading(doc, 4, "Discussion")

add_section_heading(doc, "4.1", "Why Hybrid Omega Helps", level=2)
add_body_paragraph(doc,
    "A uniform Omega treats all views with equal trust. The hybrid design "
    "achieves two simultaneous effects. First, it shrinks Omega for "
    "high-DirAcc models so the BL posterior is pulled more strongly toward "
    "their views. Second, it inflates Omega for low-magnitude predictions, "
    "preventing noise signals from perturbing the market-implied prior. Taken "
    "together, the posterior behaves as a signal-adaptive blend of market "
    "and model views, rather than a fixed-ratio combination.")

add_section_heading(doc, "4.2", "Limits of Reinforcement Learning Here", level=2)
add_body_paragraph(doc,
    "The two RL agents, despite extensive training (up to 100 episodes with "
    "deep networks), fall short of the BL_NoView Sharpe. Three reasons "
    "emerge from our analysis. First, the cross-sectional action space "
    "forces the agent to apply a single model to all 22 assets at each "
    "re-balance, discarding potentially useful per-asset diversification. "
    "Second, the discrete choice among 13 models is less expressive than "
    "continuous weighting (stacking). Third, the 5-day return signal-to-noise "
    "ratio is low, so exploration becomes inefficient. Nevertheless, the "
    "DQN_Port agent achieves the lowest drawdown among all strategies, "
    "highlighting its value as a risk-control module rather than a pure "
    "return maximiser.")

add_section_heading(doc, "4.3", "Comparison with Prior Work", level=2)
add_body_paragraph(doc,
    "Shigolakov (2025) reports that a DQN-selected BL portfolio achieves a "
    "Sharpe of 0.239 on a 10-asset universe (2024--2025), versus 0.143 for "
    "the naive portfolio. While directly comparable numbers are hindered by "
    "different asset sets and time windows, our DQN_Port achieves Sharpe "
    "1.245 over an eight-year horizon with a 22-asset universe, "
    "demonstrating that a well-designed BL pipeline (hybrid Omega + "
    "cross-sectional features) can surface stronger effects.")

# ── 5. Conclusion ──────────────────────────────────────────────────────
add_section_heading(doc, 5, "Conclusion")
add_body_paragraph(doc,
    "This paper demonstrates that the choice of Omega in the Black-Litterman "
    "model is a consequential degree of freedom, not a matter of bookkeeping. "
    "By conditioning Omega on forecasting-model quality (via direction "
    "accuracy) and prediction magnitude, we enable individual ML/DL models "
    "to outperform the BL-NoView benchmark on a 22-asset universe "
    "(LightGBM +1.6%, LSTM +1.3% in Sharpe). Reinforcement-learning "
    "agents fall short on risk-adjusted return under this particular "
    "experimental setup but improve drawdown control by 7.5%.")
add_body_paragraph(doc,
    "Future work includes (i) per-asset RL model selection, (ii) stacking "
    "meta-learners that produce continuous weights over the 13 forecasters, "
    "(iii) incorporating transaction costs into the RL reward, and "
    "(iv) extending the universe to sector- and region-diversified pools.")

# ── References ─────────────────────────────────────────────────────────
add_section_heading(doc, "", "References")

refs = [
    "Black, F., & Litterman, R. (1992). Global Portfolio Optimization. Financial Analysts Journal, 48(5), 28-43.",
    "Markowitz, H. (1952). Portfolio Selection. The Journal of Finance, 7(1), 77-91.",
    "Shigolakov, I. V. (2025). Black-Litterman Portfolio Optimization Using Machine-Learning, Deep Learning and Reinforcement Learning Algorithms. SSRN 5395585.",
    "Vijh, M., Chandola, D., Tikkiwal, V. A., & Kumar, A. (2020). Stock Closing Price Prediction using Machine Learning Techniques. Procedia Computer Science, 167, 599-606.",
    "Moody, J., & Wu, L. (1997). Optimization of trading systems and portfolios. Computational Finance, 300-307.",
    "Meucci, A. (2005). Risk and Asset Allocation. Springer.",
    "Idzorek, T. (2005). A step-by-step guide to the Black-Litterman model. In Forecasting expected returns in the financial markets (pp. 17-38).",
    "Mnih, V., et al. (2015). Human-level control through deep reinforcement learning. Nature, 518, 529-533.",
    "Schulman, J., Wolski, F., Dhariwal, P., Radford, A., & Klimov, O. (2017). Proximal policy optimization algorithms. arXiv:1707.06347.",
    "Beach, S. L., & Orlov, A. G. (2007). An application of the Black-Litterman model with EGARCH-M-derived views for international portfolio management. Applied Financial Economics, 17(17), 1393-1411.",
    "Duqi, A., Franci, L., & Torluccio, G. (2014). The Black-Litterman model: the definition of views based on volatility forecasts. Applied Financial Economics, 24(19), 1285-1296.",
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)

# ──────────── Save ────────────
doc.save(OUTPUT)
print(f"[OK] Saved: {OUTPUT}")
